import json
from docx import Document
from docx.shared import Inches
import os


def read_json(file_path):
    with open(file_path, "r") as file:
        data = json.load(file)
    return data


def add_paragraph(document, text):
    document.add_paragraph(text)


def add_image(document, image_path, width=None):
    if os.path.exists(image_path):
        if width:
            document.add_picture(image_path, width=Inches(width))
        else:
            document.add_picture(image_path)
    else:
        document.add_paragraph(f"[Image not found: {image_path}]")


def create_docx_from_json(json_data, output_path):
    document = Document()

    # Check if the data is a string and decode it
    if isinstance(json_data, str):
        json_data = json.loads(json_data)

    # Add title
    document.add_heading(json_data["title"], level=1)

    # Iterate through the story sections
    for section in json_data["storyline"]:
        if "text" in section:
            add_paragraph(document, section["text"])
        if "image" in section:
            add_image(document, section["image"], width=6)
        if "description" in section:
            add_paragraph(document, section["description"])
        if "characters" in section:
            for character in section["characters"]:
                add_paragraph(document, f"Occupation: {character['occupation']}")
                add_paragraph(document, character["description"])
                add_image(document, character["image"], width=2)
        if "clue" in section:
            add_paragraph(document, f"Clue: {section['clue']}")
            add_paragraph(document, section["description"])
            add_image(document, section["image"], width=4)
            add_image(document, section["cryptogram_image"], width=2)
        if "solution" in section:
            add_paragraph(document, section["solution"])
            add_image(document, section["image"], width=6)
            add_paragraph(document, section["description"])

    document.save(output_path)
    print(f"Document saved to {output_path}")
    return "Document saved successfully.{output_path}"


# # Paths
# json_file_path = "generated_story.json"
# output_docx_path = "generated_story.docx"

# #  Read JSON data
# json_data = read_json(json_file_path)

# # # Create DOCX file from JSON data
# create_docx_from_json(json_data, output_docx_path)
